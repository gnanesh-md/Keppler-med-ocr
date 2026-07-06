# modules/pdf_summarizer.py  — Medical Case Summary Engine v1.0
"""
Reads a multi-page hospital case file PDF (images or text),
sends each page through the vision model (qwen2.5vl:32b),
builds a structured clinical summary, and exports it as a
professionally styled 4-5 page PDF, DOCX, or Markdown file.
"""
import fitz                          # PyMuPDF
from openai import OpenAI
import base64
import io
import re
import time
import json
import markdown as md_lib
from PIL import Image
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import concurrent.futures
# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────
OCR_MODEL     = "qwen2.5-vl-7b"   # vision model for reading scanned pages
SUMMARY_MODEL = "qwen2.5-vl-7b"   # text model for summarising
MAX_IMG_DIM   = 512              # Reduced resolution for faster extraction
JPEG_QUALITY  = 75

client = OpenAI(base_url="http://localhost:8700/v1", api_key="EMPTY")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 — Extract text from each PDF page using vision model
# ─────────────────────────────────────────────────────────────────────────────

def _page_to_image_bytes(page: fitz.Page) -> bytes:
    mat = fitz.Matrix(1.5, 1.5)           # 1.5x zoom = ~108 DPI effective
    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    w, h = img.size
    if max(w, h) > MAX_IMG_DIM:
        scale = MAX_IMG_DIM / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)), Image.BILINEAR)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=JPEG_QUALITY)
    return buf.getvalue()


def extract_page_text(page_img_bytes: bytes, page_num: int) -> str:
    prompt = (
        f"This is page {page_num} of a hospital medical case file. "
        "Extract ALL text exactly as written including handwritten notes, "
        "table values, checkboxes (mark as [x] or [ ]), dates, vitals, "
        "drug names and doses. Output plain text only, no commentary."
    )
    try:
        base64_img = base64.b64encode(page_img_bytes).decode('utf-8')
        resp = client.chat.completions.create(
            model=OCR_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_img}"}}
                ]
            }],
            temperature=0,
            max_tokens=1024,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"[Page {page_num} extraction failed: {e}]"


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2 — Enterprise Summarization Pipeline
# ─────────────────────────────────────────────────────────────────────────────

# (Replaced old generate_structured_summary with modules.summarizer package)



# ─────────────────────────────────────────────────────────────────────────────
# STEP 3 — Export to PDF (styled, 4-5 pages)
# ─────────────────────────────────────────────────────────────────────────────

PDF_CSS = """
@page {
    size: A4 portrait;
    margin: 1.8cm 1.5cm 1.8cm 1.5cm;
}
body {
    font-family: Arial, Helvetica, sans-serif;
    font-size: 10px;
    color: #222;
    line-height: 1.5;
}
.cover {
    text-align: center;
    padding: 30px 0 20px;
    border-bottom: 3px solid #1a5276;
    margin-bottom: 20px;
}
.cover h1 {
    font-size: 18px;
    color: #1a5276;
    margin: 0 0 4px;
    letter-spacing: 1px;
}
.cover h2 {
    font-size: 13px;
    color: #444;
    margin: 0 0 16px;
    font-weight: normal;
}
.cover .meta-box {
    display: inline-block;
    background: #eaf2ff;
    border: 1px solid #1a5276;
    border-radius: 6px;
    padding: 10px 30px;
    font-size: 11px;
    color: #1a5276;
    font-weight: bold;
}
h1 { font-size: 16px; color: #1a5276; border-bottom: 2px solid #1a5276;
     padding-bottom: 3px; margin-top: 18px; margin-bottom: 6px; }
h2 { font-size: 13px; color: #1a5276; background: #eaf2ff;
     padding: 4px 8px; border-left: 4px solid #1a5276;
     margin-top: 14px; margin-bottom: 6px; }
h3 { font-size: 11px; color: #1a5276; margin-top: 10px; margin-bottom: 4px; }
table {
    width: 100%;
    border-collapse: collapse;
    margin: 8px 0 12px;
    font-size: 9.5px;
    page-break-inside: avoid;
}
th {
    background: #1a5276;
    color: white;
    padding: 5px 7px;
    text-align: left;
    font-weight: bold;
}
td {
    border: 1px solid #bbb;
    padding: 4px 7px;
    vertical-align: top;
}
tr:nth-child(even) td { background: #f4f8ff; }
p { margin: 3px 0; }
ul { margin: 3px 0 6px 16px; padding: 0; }
li { margin: 1px 0; }
strong { color: #1a5276; }
.section-divider {
    border: none;
    border-top: 1px dashed #aaa;
    margin: 10px 0;
}
.footer-note {
    font-size: 8px;
    color: #888;
    text-align: center;
    margin-top: 20px;
    border-top: 1px solid #ddd;
    padding-top: 6px;
}
"""


def _md_to_html(summary_md: str, patient_name: str, ip_no: str, doctor: str, nurse: str) -> str:
    safe = re.sub(r'<\|.*?\|>', '', summary_md)
    body_html = md_lib.markdown(safe, extensions=['tables', 'nl2br'])

    cover = f"""
    <div class="cover">
        <h1>GREAT EASTERN MEDICAL SCHOOL &amp; HOSPITAL</h1>
        <h2>Promoted by Aditya Educational Society</h2>
        <div class="meta-box">
            PATIENT CASE SUMMARY<br/>
            <span style="font-size:13px">{patient_name}</span><br/>
            IP No: {ip_no} &nbsp;|&nbsp; Doctor: {doctor} &nbsp;|&nbsp; Nurse: {nurse}
        </div>
        <p style="font-size:9px;color:#888;margin-top:12px">
            Generated on {time.strftime('%d %B %Y at %I:%M %p')} &nbsp;|&nbsp; Confidential — For Medical Records Only
        </p>
    </div>
    """

    footer = """
    <div class="footer-note">
        This summary was auto-generated from scanned case file documents using Keppler AI (OCR + LLM).
        Verify all clinical data against original records before use.
    </div>
    """

    return f"""<html>
<head><meta charset="utf-8"/>
<style>{PDF_CSS}</style>
</head>
<body>
{cover}
{body_html}
{footer}
</body></html>"""


def generate_summary_pdf(summary_md: str, patient_name: str, ip_no: str, doctor: str, nurse: str) -> bytes | str:
    try:
        html = _md_to_html(summary_md, patient_name, ip_no, doctor, nurse)
        out = io.BytesIO()
        status = pisa.CreatePDF(io.StringIO(html), dest=out, encoding="utf-8")
        return out.getvalue() if not status.err else f"PDF Error: {status.err}"
    except Exception as e:
        return f"PDF Error: {e}"


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4 — Export to DOCX (styled Word document)
# ─────────────────────────────────────────────────────────────────────────────

def generate_summary_docx(summary_md: str, patient_name: str, ip_no: str, doctor: str, nurse: str) -> bytes | str:
    try:
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        # Cover heading
        title = doc.add_heading("GREAT EASTERN MEDICAL SCHOOL & HOSPITAL", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

        sub = doc.add_paragraph("PATIENT CASE SUMMARY")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].bold = True
        sub.runs[0].font.size = Pt(13)

        info = doc.add_paragraph(f"Patient: {patient_name}   |   IP No: {ip_no}")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.runs[0].font.size = Pt(10)

        staff = doc.add_paragraph(f"Doctor: {doctor}   |   Nurse: {nurse}")
        staff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        staff.runs[0].font.size = Pt(9)

        date_p = doc.add_paragraph(
            f"Generated: {time.strftime('%d %B %Y at %I:%M %p')} | Confidential"
        )
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.runs[0].font.size = Pt(8)
        date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        # Parse and write markdown sections
        clean = re.sub(r'<\|.*?\|>', '', summary_md)
        in_table = False
        table_obj = None

        for line in clean.split('\n'):
            line = line.rstrip()

            if line.startswith('## '):
                in_table = False
                h = doc.add_heading(line[3:], level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

            elif line.startswith('### '):
                in_table = False
                h = doc.add_heading(line[4:], level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

            elif line.startswith('- ') or line.startswith('* '):
                in_table = False
                p = doc.add_paragraph(style='List Bullet')
                text = line[2:].strip()
                # Handle **bold** inline
                parts = re.split(r'\*\*(.+?)\*\*', text)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True

            elif line.startswith('|'):
                if '---' in line:
                    continue
                cells = [c.strip() for c in line.strip('|').split('|')]
                cells = [c for c in cells if c]
                if not cells:
                    continue
                if not in_table:
                    table_obj = doc.add_table(rows=1, cols=len(cells))
                    table_obj.style = 'Table Grid'
                    hdr = table_obj.rows[0]
                    for i, c in enumerate(cells):
                        cell = hdr.cells[i]
                        cell.text = c
                        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(c)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        cell._tc.get_or_add_tcPr()
                    in_table = True
                else:
                    row = table_obj.add_row()
                    for i, c in enumerate(cells):
                        if i < len(row.cells):
                            row.cells[i].text = c

            elif line.strip() == '' or line.startswith('---'):
                in_table = False
                if line.strip() == '':
                    doc.add_paragraph()

            else:
                in_table = False
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
                    run.font.size = Pt(10)

        # Footer note
        doc.add_paragraph()
        note = doc.add_paragraph(
            "⚠ This summary was auto-generated by Keppler AI from scanned case file documents. "
            "Verify all clinical data against original records before use."
        )
        note.runs[0].font.size = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        return f"DOCX Error: {e}"


# ─────────────────────────────────────────────────────────────────────────────
# MAIN STREAMLIT UI
# ─────────────────────────────────────────────────────────────────────────────

def _extract_patient_meta(summary_md: str) -> tuple[str, str, str, str]:
    """Try to pull patient name, IP number, doctor, and nurse from the summary text."""
    name = "Patient"
    ip   = "—"
    doctor = "—"
    nurse = "—"
    name_m = re.search(r'(?:name|patient)[:\s*_]+([A-Z][A-Z\s]+)', summary_md, re.I)
    ip_m   = re.search(r'IP\s*(?:No|Number)?[:\s._]+([A-Z0-9]+)', summary_md, re.I)
    doc_m  = re.search(r'(?:Treating Doctor|Consultant)[:\s*_]+([A-Z][A-Za-z\s.]+)', summary_md, re.I)
    nurse_m= re.search(r'(?:Treating Nurse|Staff)[:\s*_]+([A-Z][A-Za-z\s.]+)', summary_md, re.I)
    if name_m:
        name = name_m.group(1).strip().title()
    if ip_m:
        ip = ip_m.group(1).strip()
    if doc_m:
        doctor = doc_m.group(1).strip().title()
    if nurse_m:
        nurse = nurse_m.group(1).strip().title()
    return name, ip, doctor, nurse


def run_summary_pipeline(pdf_bytes: bytes, filename: str = "document.pdf", progress_cb=None, clear_cache: bool = False) -> dict:
    """
    Streamlit-free pipeline: OCR each page of a case-file PDF concurrently, then
    map-reduce summarize into a structured clinical report.

    progress_cb, if given, is called with a float in [0, 1] — 0.0-0.5 covers
    per-page OCR, 0.5-1.0 covers chunk summarization.

    Returns {"summary_md": str, "page_texts": dict[int, str], "patient_meta": (name, ip, doctor, nurse)}.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)

    images_to_process = {}
    for i in range(total_pages):
        page = doc[i]
        images_to_process[i + 1] = _page_to_image_bytes(page)
    doc.close()

    page_texts = {}
    completed = 0
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        futures = {
            executor.submit(extract_page_text, img_bytes, pg_num): pg_num
            for pg_num, img_bytes in images_to_process.items()
        }
        for future in concurrent.futures.as_completed(futures):
            pg_num = futures[future]
            page_texts[pg_num] = future.result()
            completed += 1
            if progress_cb:
                progress_cb(0.5 * completed / total_pages)

    page_strings = [page_texts[pg] for pg in sorted(page_texts.keys())]

    from modules.summarizer.chunker import TextChunker
    from modules.summarizer.hierarchical import HierarchicalSummarizer
    from modules.summarizer.aggregator import MasterAggregator

    chunker = TextChunker(pages=page_strings, chunk_size=1, document_id=filename)
    if clear_cache:
        chunker.clear_cache()

    hierarchical = HierarchicalSummarizer(model_name=SUMMARY_MODEL)

    def _update_progress(current, total):
        if progress_cb:
            progress_cb(0.5 + 0.4 * current / total)

    master_data = hierarchical.process_chunks(chunker, progress_callback=_update_progress)

    overall_summary = hierarchical.generate_overall_summary(master_data.get("page_summaries", []))
    master_data["overall_summary"] = overall_summary

    summary_md = MasterAggregator.build_markdown_report(master_data)

    if progress_cb:
        progress_cb(1.0)

    patient_meta = _extract_patient_meta(summary_md)

    return {
        "summary_md": summary_md,
        "page_texts": page_texts,
        "patient_meta": patient_meta,
    }
