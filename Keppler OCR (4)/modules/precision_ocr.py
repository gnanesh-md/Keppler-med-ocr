# modules/precision_ocr.py  — Production OCR Engine v2.0
"""
Production-level OCR pipeline using gemma4:26b vision model.
Features:
  - 5 adaptive image preprocessing strategies with auto-retry
  - Multi-page PDF support (processes all pages, merges results)
  - Image upscaling for low-resolution scans
  - Robust response parsing with fallback chain
  - Per-step error surfacing (no silent failures)
  - Archive vault integration
  - Download in MD / PDF / DOCX formats
"""
import time
import io
import asyncio
import json
import logging
import fitz          # PyMuPDF
import re
import math
import numpy as np
import cv2
from PIL import Image, ImageEnhance, ImageOps, ImageFilter
from openai import OpenAI
import base64
import markdown
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import pandas as pd
from core.config import settings
from modules.unified_resolver import resolve_entities_in_text
from modules.layout_detector import DocLayoutDetector
from modules.reading_order import ReadingOrderEngine
from modules.region_ocr import AsyncRegionOCR
from modules.table_extractor import TableExtractor
from modules.medical_corrector import MedicalCorrector
from modules.confidence_engine import ConfidenceEngine

# Bounded timeout (SDK default is 600s) — a hung/overloaded vLLM server should
# fail one call fast and let Celery's retry policy handle it, not silently
# block a worker for up to 10 minutes per call.
# Always settings.VLLM_BASE_URL, never a hardcoded "localhost:8700" — a
# hardcoded host string here silently breaks in any containerized deployment.
client = OpenAI(base_url=settings.VLLM_BASE_URL, api_key="EMPTY", timeout=90.0)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 0.  PAGE ORIENTATION DETECTION (0/90/180/270)
# ---------------------------------------------------------------------------
def detect_page_orientation(img: Image.Image) -> int:
    """
    Detects the rotation (0/90/180/270, clockwise) needed to make a page
    upright, run once per page before the fine-angle deskew strategies below
    (which only correct small ±0.5-15 degree skew, not full rotations).

    Two-step, since a single metric can't do this alone:
      1. Row-projection-variance picks the axis (portrait {0,180} vs
         landscape {90,270}) — properly oriented text has strong periodic
         horizontal line-banding that a 90-degree rotation destroys.
      2. Row-projection variance looks nearly identical for 0 vs 180 (upside
         -down text has the same line-banding pattern), so those two are
         disambiguated with a real OCR-confidence probe on a center crop —
         whichever orientation the model reads with higher confidence wins.
         This deliberately uses the primary vLLM engine (not TrOCR): tested
         against a real scanned page, TrOCR's confidence didn't track actual
         orientation at all (0.34 for the correct upright crop vs 0.37 for
         the wrong upside-down one — noise), whereas vLLM scored the correct
         orientation clearly higher (0.68 vs 0.42) and produced coherent text
         only for the right one. Costs two small vLLM calls per page.

    The confidence probe is a single noisy model signal, not a certainty. For
    the 0-vs-180 case it was reliable across every real document tested (a
    consistent, clear confidence margin favoring the correct upright read).
    For the 90-vs-270 case it was NOT: tested against two different real
    documents it confidently favored the *wrong* member of the pair both
    times (not just a low-margin coin flip the margin gate below would
    catch) — the direction of a quarter-turn is apparently harder for this
    signal to judge reliably than a half-turn is. Rather than ship something
    that can confidently flip an already-correct page 90 degrees the wrong
    way (actively worse than the pre-Phase-2 baseline of not rotating at
    all), quarter-turn axis pages are intentionally left uncorrected here —
    flagged via a log line for visibility, not silently guessed. Only the
    half-turn (0/180) correction is applied automatically. Reliably picking
    quarter-turn direction is left as follow-up work (e.g. an actual OSD
    model) rather than shipped as a guess.
    """
    arr = np.array(img.convert("L"))

    def _row_projection_variance(rotation_deg: int) -> float:
        rotated = np.rot90(arr, k=rotation_deg // 90)
        return float(np.var(rotated.sum(axis=1)))

    portrait_axis = _row_projection_variance(0) >= _row_projection_variance(90)
    if not portrait_axis:
        logger.info(
            "Page appears rotated a quarter-turn (90/270) but the direction "
            "isn't reliably determined by the current confidence probe — "
            "leaving uncorrected rather than guessing."
        )
        return 0
    candidates = (0, 180)

    w, h = img.size
    crop_w, crop_h = min(w, 700), min(h, 300)
    cx, cy = w // 2, h // 2
    center_crop = img.crop(
        (cx - crop_w // 2, cy - crop_h // 2, cx + crop_w // 2, cy + crop_h // 2)
    )

    scores = {}
    for rot in candidates:
        test_img = center_crop.rotate(-rot, expand=True) if rot else center_crop
        scores[rot] = _vllm_orientation_confidence(test_img)

    ranked = sorted(candidates, key=lambda r: scores[r], reverse=True)
    winner, runner_up = ranked[0], ranked[1]
    MIN_ABSOLUTE_CONFIDENCE = 0.5
    MIN_MARGIN_RATIO = 1.3  # winner must be a clearly better read, not a coin flip
    if scores[winner] < MIN_ABSOLUTE_CONFIDENCE or scores[winner] < MIN_MARGIN_RATIO * max(scores[runner_up], 1e-6):
        return 0
    return winner


def _vllm_orientation_confidence(img: Image.Image) -> float:
    """Mean per-token log-prob (exponentiated) of a short vLLM read of `img`,
    used only to compare orientation candidates against each other.

    Deliberately creates its own client per call rather than using the
    module-level `client` above: that client is constructed at import time,
    and Celery's prefork worker model forks child processes *after* the
    module (and its at-import-time HTTP client/connection pool) is loaded —
    reusing a pre-fork httpx client from a forked child is a known way to
    deadlock silently. Confirmed by reproduction: page tasks hung for
    ~10 minutes (the OpenAI SDK's default request timeout) with zero
    progress the first time this code path ran under the real Celery worker,
    vanishing once each call got its own client instead of sharing the
    pre-fork one.
    """
    try:
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        local_client = OpenAI(base_url=settings.VLLM_BASE_URL, api_key="EMPTY", timeout=30.0)
        resp = local_client.chat.completions.create(
            model="qwen2.5-vl-7b",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": "Read the text in this image exactly as written."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                ],
            }],
            temperature=0,
            max_tokens=40,
            logprobs=True,
            top_logprobs=1,
        )
        content = resp.choices[0].logprobs.content if resp.choices[0].logprobs else None
        if not content:
            return 0.0
        return float(np.exp(np.mean([t.logprob for t in content])))
    except Exception as e:
        logger.warning(f"Orientation confidence probe failed: {e}")
        return 0.0


def apply_orientation(img: Image.Image) -> Image.Image:
    """Detects and corrects full-rotation page orientation. No-op (cheap) for
    the common already-upright case; returns the input unchanged then."""
    rotation = detect_page_orientation(img)
    if rotation == 0:
        return img
    logger.info(f"Correcting page rotation by {rotation} degrees")
    return img.rotate(-rotation, expand=True)


# ---------------------------------------------------------------------------
# 0b. BLANK PAGE DETECTION
# ---------------------------------------------------------------------------
def is_blank_page(img: Image.Image, ink_ratio_threshold: float = 0.01) -> bool:
    """
    Cheap (no model call) check for near-empty separator/blank pages, common
    in real hospital scan batches. Thresholds against the page's own median
    background level (works for off-white/gray scans, not just pure white)
    and measures what fraction of pixels are meaningfully darker than that.

    Tuned against real documents in uploads/: genuine content pages measured
    6-17% ink coverage; a synthetic blank page (with scanner noise) measured
    0%, and a near-blank page with only a small stamp/mark measured 0.4% —
    both comfortably under this threshold, with a wide margin to real content.
    """
    gray = np.array(img.convert("L"), dtype=np.float32)
    bg_level = np.median(gray)
    ink_mask = gray < (bg_level - 30)
    ink_ratio = float(ink_mask.mean())
    return ink_ratio < ink_ratio_threshold

# ---------------------------------------------------------------------------
# 1.  IMAGE PREPROCESSING STRATEGIES
MIN_DIM = 600   # minimum pixel dimension; smaller images are upscaled

def _upscale_if_small(img: Image.Image) -> Image.Image:
    """Upscale image if either dimension is below MIN_DIM."""
    w, h = img.size
    if min(w, h) < MIN_DIM:
        scale = MIN_DIM / min(w, h)
        # Cap scale factor to prevent line crops (e.g. 800x40) from becoming massively oversized
        if scale > 2.5:
            scale = 2.5
        # Ensure the maximum dimension doesn't exceed 3000 after upscaling
        if max(w, h) * scale > 3000:
            scale = 3000 / max(w, h)
        
        if scale > 1.0:
            new_w, new_h = int(w * scale), int(h * scale)
            img = img.resize((new_w, new_h), Image.BILINEAR)
    return img

def strategy_original(img: Image.Image) -> Image.Image:
    """Strategy 0 — Send the image with minimal changes (fix orientation only)."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    return img.convert("RGB")

def strategy_mild_enhance(img: Image.Image) -> Image.Image:
    """Strategy 1 — Mild sharpening + contrast boost. Best for printed forms."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    img = img.convert("RGB")
    img = ImageEnhance.Sharpness(img).enhance(1.8) 
    img = ImageEnhance.Contrast(img).enhance(1.5)
    return img

def strategy_grayscale_boost(img: Image.Image) -> Image.Image:
    """Strategy 2 — Grayscale + moderate contrast. Good for photocopies."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.8)
    gray = ImageEnhance.Sharpness(gray).enhance(2.0)
    return gray.convert("RGB")

def strategy_adaptive_threshold(img: Image.Image) -> Image.Image:
    """Strategy 3 — Near-binarization for faded handwriting on white background."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    gray = img.convert("L")
    # Auto-level: stretch histogram
    arr = np.array(gray, dtype=np.float32)
    p2, p98 = np.percentile(arr, 2), np.percentile(arr, 98)
    if p98 > p2:
        arr = np.clip((arr - p2) / (p98 - p2) * 255, 0, 255)
    gray = Image.fromarray(arr.astype(np.uint8))
    gray = ImageEnhance.Contrast(gray).enhance(2.2)
    # Unsharp mask to make thin strokes pop
    gray = gray.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
    return gray.convert("RGB")

def strategy_denoised(img: Image.Image) -> Image.Image:
    """Strategy 4 — Denoise first then enhance. For camera photos of documents."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    img = img.convert("RGB")
    # Median filter removes camera noise
    img = img.filter(ImageFilter.MedianFilter(size=3))
    img = ImageEnhance.Contrast(img).enhance(1.6)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    return img

def _deskew_cv(img_bgr: np.ndarray) -> np.ndarray:
    """Hough-line-based skew detection + rotation correction."""
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(edges, 1, np.pi / 180, 100, minLineLength=100, maxLineGap=10)
    if lines is None:
        return img_bgr
    # reshape(-1, 4), not lines[:, 0]: HoughLinesP's output shape is (N, 1, 4)
    # on some opencv-python-headless versions and (N, 4) on others (hit this
    # for real — a fresh `pip install` picked up a newer version than the
    # long-running host venv had, and lines[:, 0] silently returned scalars
    # instead of (x1,y1,x2,y2) tuples on the (N, 4) shape, breaking deskew
    # for every image). reshape(-1, 4) is correct for both.
    angles = [np.degrees(np.arctan2(y2 - y1, x2 - x1)) for (x1, y1, x2, y2) in lines.reshape(-1, 4)]
    if not angles:
        return img_bgr
    median_angle = np.median(angles)
    # Only correct noticeable skew; skip near-0 (already straight) and near-90 (rotated page, different problem).
    if not (0.5 < abs(median_angle) < 15):
        return img_bgr
    h, w = img_bgr.shape[:2]
    m = cv2.getRotationMatrix2D((w // 2, h // 2), median_angle, 1.0)
    return cv2.warpAffine(img_bgr, m, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)


def strategy_deskew_clahe_denoise(img: Image.Image) -> Image.Image:
    """Strategy 5 — Deskew + CLAHE contrast + non-local-means denoise. For crooked
    camera photos / fax-quality scans where the other strategies leave visible tilt
    or uneven lighting."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    bgr = cv2.cvtColor(np.array(img.convert("RGB")), cv2.COLOR_RGB2BGR)
    bgr = _deskew_cv(bgr)

    lab = cv2.cvtColor(bgr, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    l = clahe.apply(l)
    bgr = cv2.cvtColor(cv2.merge((l, a, b)), cv2.COLOR_LAB2BGR)
    bgr = cv2.fastNlMeansDenoisingColored(bgr, None, 10, 10, 7, 21)

    return Image.fromarray(cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB))


STRATEGIES = [
    ("Original",              strategy_original),
    ("Mild Enhancement",      strategy_mild_enhance),
    ("Grayscale Boost",       strategy_grayscale_boost),
    ("Adaptive Threshold",    strategy_adaptive_threshold),
    ("Denoise + Enhance",     strategy_denoised),
    ("Deskew + CLAHE",        strategy_deskew_clahe_denoise),
]

def img_to_bytes(img: Image.Image, quality: int = 92) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# 2.  RESPONSE CLEANING
# ---------------------------------------------------------------------------
def clean_output(raw: str) -> str:
    """
    Remove model control tokens and code fences.
    Falls back to the raw text if cleaning removes everything.
    Preserves medical symbols like < > in reference ranges.
    """
    if not raw:
        return ""

    text = raw

    # Strip content before <|text|> markers (some model variants)
    if "<|text|>" in text:
        text = text.split("<|text|>")[-1]
    if "</|text|>" in text:
        text = text.split("</|text|>")[0]

    # Remove thinking blocks — both formats used by different model variants
    text = re.sub(r'<\|think\|>.*?</\|think\|>', '', text, flags=re.DOTALL)   # pipe format
    text = re.sub(r'<think>.*?</think>',           '', text, flags=re.DOTALL)   # plain format

    # Remove code fences
    text = re.sub(r'```[a-z]*\n?', '', text)
    text = text.replace('```', '')

    cleaned = text.strip()

    # Safety fallback: if stripping removed EVERYTHING, return raw minus fences
    if not cleaned:
        fallback = re.sub(r'```[a-z]*\n?', '', raw).replace('```', '').strip()
        return fallback

    return cleaned

# ---------------------------------------------------------------------------
# 3.  PDF MULTI-PAGE LOADER
# ---------------------------------------------------------------------------
def load_pdf_page_count(file_bytes: bytes) -> int:
    """Page count without rendering — cheap, used to size the Celery chord fan-out."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    count = len(doc)
    doc.close()
    return count


def load_pdf_page(file_bytes: bytes, page_idx: int) -> Image.Image:
    """Render a single PDF page as a PIL image at 300 DPI. Used by the per-page
    Celery task so a worker only ever holds one rendered page in memory at a
    time, regardless of document length (contrast with load_pdf_pages below,
    which is a convenience for small/synchronous callers only)."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc.load_page(page_idx)
    pix = page.get_pixmap(dpi=300)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    img.load()  # decode before closing the underlying doc/buffer
    doc.close()
    return img


def load_pdf_page_thumbnail(file_bytes: bytes, page_idx: int, dpi: int = 72) -> Image.Image:
    """Cheap low-DPI render of a single page, used only for perceptual-hash
    duplicate detection at dispatch time (see workers/celery_app.py) — avoids
    paying full 300 DPI render cost twice (once to hash, once to OCR) for
    every page just to detect the minority that are duplicates."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc.load_page(page_idx)
    pix = page.get_pixmap(dpi=dpi)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    img.load()
    doc.close()
    return img


def page_phash(img: Image.Image) -> str:
    """Perceptual hash for near-duplicate page detection (tolerant to minor
    re-scan/re-compress differences that an exact byte/MD5 hash would miss)."""
    import imagehash
    return str(imagehash.phash(img))


def load_pdf_pages(file_bytes: bytes, max_pages: int | None = None) -> list[Image.Image]:
    """Extract pages from a PDF as PIL images at 300 DPI (Balanced for Qwen).

    Loads every requested page into memory at once — fine for the synchronous/
    single-process callers (tests, small documents) but NOT what the Celery OCR
    pipeline uses for large PDFs; that path calls load_pdf_page_count/
    load_pdf_page one page at a time instead. No page cap by default."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total = len(doc) if max_pages is None else min(len(doc), max_pages)
    pages = []
    for i in range(total):
        page = doc.load_page(i)
        pix  = page.get_pixmap(dpi=300)
        pages.append(Image.open(io.BytesIO(pix.tobytes("png"))))
    doc.close()
    return pages

# ---------------------------------------------------------------------------
# 5.  EXPORT GENERATORS  (PDF / DOCX)
# ---------------------------------------------------------------------------
def _is_table_separator(line: str) -> bool:
    """
    True for a Markdown table separator/alignment row (e.g. '|---|---|',
    '|:--:|-|'). Shared by every download format (PDF/DOCX/Excel/JSON) so a
    row is classified the same way everywhere — before this was unified,
    DOCX used a bare `'---' in line` check (misses a minimal '|-|-|-|'
    separator, and can wrongly swallow a real data row that merely contains
    the substring '---') while Excel/JSON didn't allow ':' (misreads a
    colon-alignment separator like '|:---:|' as a real data row).
    """
    return bool(re.fullmatch(r'[\s|:\-]+', line)) and '-' in line


def _normalize_markdown_tables(text: str) -> str:
    """
    Rebuilds every block of consecutive '|'-led lines into a strict, valid
    Markdown table (matching column counts, a proper '|---|' separator row)
    before handing text to python-markdown's 'tables' extension.

    That extension is strict: a separator row whose dash-group count doesn't
    match the header, or a data row with more/fewer cells than the header,
    makes it silently bail and render the *entire* block as one plain
    paragraph of literal pipe characters instead of a table. Real vision-
    model OCR output hits this constantly (an extra/missing cell on one row,
    a mangled separator), so without this repair pass the PDF export renders
    those tables as an unstructured wall of text. generate_docx/_excel/_json
    sidestep this by parsing pipe-led lines manually instead of relying on
    the strict extension; this brings the PDF path in line with them.
    """
    lines = text.split('\n')
    out = []
    i, n = 0, len(lines)
    while i < n:
        if lines[i].strip().startswith('|'):
            j = i
            block = []
            while j < n and lines[j].strip().startswith('|'):
                block.append(lines[j].strip())
                j += 1

            rows = []
            for raw in block:
                if _is_table_separator(raw):
                    continue  # drop the model's own (possibly malformed) separator row
                rows.append([c.strip() for c in raw.strip('|').split('|')])

            if rows:
                col_count = len(rows[0])
                fixed = []
                for r in rows:
                    if len(r) < col_count:
                        r = r + [''] * (col_count - len(r))
                    elif len(r) > col_count:
                        r = r[:col_count]
                    # xhtml2pdf collapses a column's width to ~0 (ignoring any
                    # explicit/fixed width) if any cell in it is truly empty,
                    # so a blank field — which the LDSL/Healmax blueprints
                    # explicitly produce for missing data — must never reach
                    # the PDF renderer as an empty cell.
                    r = [c if c.strip() else '&nbsp;' for c in r]
                    fixed.append(r)
                out.append('| ' + ' | '.join(fixed[0]) + ' |')
                out.append('|' + '|'.join(['---'] * col_count) + '|')
                out.extend('| ' + ' | '.join(r) + ' |' for r in fixed[1:])
            i = j
        else:
            out.append(lines[i])
            i += 1
    return '\n'.join(out)


def _widest_table_column_count(text: str) -> int:
    """
    Max column count across every Markdown table row in the text (post
    _normalize_markdown_tables, so rows are already well-formed). Landscape
    orientation should key off this, not merely "does a table exist" — a
    narrow 4-column LDSL table fits comfortably in portrait, and forcing it
    into landscape actively hurts: A4 landscape has *less* vertical room
    (595pt) than portrait (842pt), so the same row count that fit on one
    portrait page spills onto a second landscape page for no benefit. A wide
    9-column Healmax table genuinely needs the extra horizontal room though.
    """
    max_cols = 0
    for line in text.split('\n'):
        line = line.strip()
        if line.startswith('|') and not _is_table_separator(line):
            max_cols = max(max_cols, len(line.strip('|').split('|')))
    return max_cols


_META_FIELD_RE = re.compile(r'^\*{0,2}([^*:|]+):\*{0,2}\s*(.*)$')


def _is_meta_field_segment(segment: str) -> bool:
    return bool(_META_FIELD_RE.match(segment.strip()))


def _render_meta_group(lines: list[str]) -> str:
    """
    Renders one or more consecutive 'Label: Value | Label: Value' lines
    (patient info, doctor/date fields, footer fields — bold or not) as a
    single borderless table with one row per line, instead of plain run-on
    text with literal '|' characters. A real HTML table is used (rather than
    e.g. flexbox) because xhtml2pdf only supports a small, table/block-based
    subset of CSS2.1 — flex layout is silently dropped.

    The divider rule goes on the LAST row's cells only, never on the <table>
    or a wrapping <div>: xhtml2pdf repeats a border set on a multi-child
    block container onto every child row instead of drawing it once at the
    bottom (verified empirically), so putting it on the container would draw
    a line under every single field row instead of one line under the group.
    (A leading rule was tried for footer groups too, to box them like a
    signature bar, but immediately following another group's own trailing
    rule it produced two near-touching lines — a trailing-only rule reads
    just as clearly as a boxed footer without that artifact.)
    """
    parsed = [[s.strip() for s in line.split('|')] for line in lines]
    max_cols = max(len(segs) for segs in parsed)
    rows_html = []
    for idx, segments in enumerate(parsed):
        cells = []
        for i, seg in enumerate(segments):
            m = _META_FIELD_RE.match(seg)
            label, value = m.group(1).strip(), m.group(2).strip()
            # A leading bold tag like "**Footer:**" is sometimes just a section
            # label prefixed onto the first real field rather than a field of
            # its own (e.g. "**Footer:** Collection Executive: John") — when
            # the value itself parses as another Label: Value pair, prefer
            # that inner pair so "Collection Executive" gets its own clean
            # field instead of being nested as a value under "Footer".
            inner = _META_FIELD_RE.match(value)
            if inner and inner.group(1).strip():
                label, value = inner.group(1).strip(), inner.group(2).strip()
            attrs = ' style="width: 40%"' if i == 0 and len(segments) > 1 else ''
            # A row with fewer fields than the widest row in this group (e.g.
            # a lone "Ultrasound Findings: N/A" alongside two-field rows)
            # needs colspan on its last cell so its border — if this ends up
            # being the group's last row — stretches the full row width
            # instead of stopping at the first column.
            if i == len(segments) - 1 and len(segments) < max_cols:
                attrs += f' colspan="{max_cols - len(segments) + 1}"'
            cells.append(f'<td class="meta-cell"{attrs}><span class="meta-label">{label}:</span> {value}</td>')
        cls = ' class="meta-last"' if idx == len(parsed) - 1 else ''
        rows_html.append(f'<tr{cls}>{"".join(cells)}</tr>')
    return f'<table class="meta-table">{"".join(rows_html)}</table>'


def _is_section_heading(line: str) -> bool:
    s = line.strip()
    return s.startswith('**') and s.endswith('**') and s.count('**') == 2 and ':' not in s[2:-2] and '|' not in s


def _preprocess_pdf_text(text: str) -> tuple[str, dict[str, str]]:
    """
    Pulls patient-info/footer 'Label: Value' lines and standalone bold
    section titles out of the normal Markdown paragraph flow and pre-renders
    them (info lines as an aligned table row, section titles as a styled
    heading) so the PDF shows a real form layout instead of a single run-on
    line of text with literal pipe characters. Each becomes a placeholder
    token here, substituted back into the real HTML after markdown.markdown()
    runs, since injecting raw HTML mid-stream would otherwise fight the
    Markdown parser's own paragraph/emphasis handling.
    """
    placeholders: dict[str, str] = {}
    out_lines: list[str] = []
    meta_buffer: list[str] = []

    def flush_meta():
        if meta_buffer:
            key = f'KEPPLERPLACEHOLDER{len(placeholders)}X'
            placeholders[key] = _render_meta_group(list(meta_buffer))
            out_lines.append(key)
            meta_buffer.clear()

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped or stripped.startswith(('|', '#', '- ', '* ')):
            flush_meta()
            out_lines.append(line)
            continue

        if _is_section_heading(stripped):
            flush_meta()
            key = f'KEPPLERPLACEHOLDER{len(placeholders)}X'
            placeholders[key] = f'<h4 class="section-heading">{stripped.strip("*")}</h4>'
            out_lines.append(key)
            continue

        segments = [s.strip() for s in stripped.split('|')]
        if segments and all(_is_meta_field_segment(s) for s in segments):
            meta_buffer.append(stripped)
            continue

        flush_meta()
        out_lines.append(line)
    flush_meta()
    return '\n'.join(out_lines), placeholders


def generate_pro_pdf(md_text: str, client_name: str) -> bytes | str:
    try:
        safe = re.sub(r'<\|.*?\|>', '', md_text)
        safe = safe.replace('<', '&lt;').replace('>', '&gt;')
        safe = safe.replace('\u2022', '*').replace('\u2713', '[x]')
        
        is_form = "LDSL" in client_name or "Healmax" in client_name
        
        if is_form:
            safe = _normalize_markdown_tables(safe)
            safe, meta_placeholders = _preprocess_pdf_text(safe)
            subtitle = "Test Requisition Form"
        else:
            meta_placeholders = {}
            subtitle = "Extraction Report"
            # Prevent markdown from parsing lists if we want them to look like raw text lines
            safe = re.sub(r'(?m)^(\s*)[-*]\s+', r'\1&#45; ', safe)
            
        orientation = "landscape" if _widest_table_column_count(safe) > 6 else "portrait"
        
        extra_css = ""
        # CSS nth-child is ignored by xhtml2pdf, so we rely on inline styles below, 
        # but keep this for standard HTML viewers if ever exported.
        if "LDSL" in client_name:
            extra_css = """
            th:nth-child(1) { width: 8%; }
            th:nth-child(2) { width: 45%; }
            th:nth-child(3) { width: 30%; }
            th:nth-child(4) { width: 17%; }
            """
        elif "Healmax" in client_name:
            extra_css = """
            th:nth-child(2) { width: 20%; }
            th:nth-child(4) { width: 20%; }
            """

        pdf_css = f"""
        @page {{ size: A4 {orientation}; margin: 1.8cm 1.5cm 1.8cm 1.5cm; }}
        body {{ font-family: Courier, monospace; font-size: 11px; color: #2a2a2a; line-height: 1.5; white-space: pre-wrap; }}
        .cover {{ padding: 10px 0 0; margin-bottom: 14px; }}
        .cover p.title {{ font-size: 22px; color: #123049; font-weight: bold; margin: 0; letter-spacing: 0.2px; font-family: Arial, sans-serif; }}
        .cover p.subtitle {{ font-size: 11px; color: #667080; margin: 3px 0 0 0; padding-bottom: 11px; border-bottom: 2px solid #d8dde2; font-family: Arial, sans-serif; }}
        h1 {{ font-size: 16px; color: #333; border-bottom: 1px solid #d8dde2; padding-bottom: 3px; margin-top: 18px; margin-bottom: 6px; font-family: Arial, sans-serif; }}
        h2 {{ font-size: 13px; color: #444; margin-top: 14px; margin-bottom: 6px; font-family: Arial, sans-serif; }}
        h3 {{ font-size: 11px; color: #555; margin-top: 10px; margin-bottom: 4px; font-family: Arial, sans-serif; }}
        table {{ width: 100%; table-layout: fixed; border-collapse: collapse; margin: 11px 0; font-size: 10px; font-family: Arial, sans-serif; }}
        thead {{ display: table-header-group; }}
        tr {{ page-break-inside: avoid; }}
        th {{ background: #eef2f7; color: #123049; padding: 8px; text-align: left; font-weight: bold; border: 1px solid #dde2e7; }}
        td {{ border: 1px solid #dde2e7; padding: 7px 8px; vertical-align: top; word-wrap: break-word; }}
        tr:nth-child(even) td {{ background: #f8fafb; }}
        p {{ margin: 4px 0; }}
        ul {{ margin: 4px 0 8px 16px; padding: 0; }}
        li {{ margin: 2px 0; }}
        strong {{ color: #123049; }}
        .meta-table {{ width: 100%; table-layout: auto; border-collapse: collapse; margin: 6px 0; font-family: Arial, sans-serif; }}
        .meta-table td.meta-cell {{ border: none; padding: 4px 10px 4px 0; font-size: 10px; }}
        .meta-table tr.meta-last td.meta-cell {{ border-bottom: 1px solid #d8dde2; padding-bottom: 9px; }}
        .meta-label {{ font-weight: bold; color: #123049; }}
        .section-heading {{ font-size: 11.5px; font-style: italic; font-weight: bold; color: #123049; margin-top: 14px; margin-bottom: 5px; border-bottom: 1px solid #d8dde2; padding-bottom: 3px; font-family: Arial, sans-serif; }}
        .footer-note {{ font-size: 9px; color: #8a93a0; text-align: center; margin-top: 16px; border-top: 1px solid #e2e6ea; padding-top: 6px; font-family: Arial, sans-serif; }}
        {extra_css}
        """

        cover = f"""
        <div class="cover">
            <p class="title">{client_name}</p>
            <p class="subtitle">{subtitle}</p>
        </div>
        """

        footer = f"""
        <div class="footer-note">
            Generated on {time.strftime('%d %B %Y at %I:%M %p')} by Keppler AI. Verify all data against original records before use.
        </div>
        """
        
        # Ensure python-markdown sees tables properly by adding a blank line before the header
        safe = re.sub(r'([^\|\n][ \t]*)\n([ \t]*\|)', r'\1\n\n\2', safe)
        
        # Let the standard python markdown parser handle the tables
        html_body = markdown.markdown(safe, extensions=['tables', 'nl2br'])

        # Substitute the pre-rendered meta-row/section-heading blocks back in,
        # then unwrap the <p>...</p> markdown wrapped around them (and drop
        # the <br> it joins adjacent placeholder lines with) — table/heading
        # markup nested inside a <p> is invalid HTML that xhtml2pdf renders
        # with stray extra spacing.
        for key, block_html in meta_placeholders.items():
            html_body = html_body.replace(key, block_html)
        block_alt = r'(?:<table class="meta-table">.*?</table>|<h4 class="section-heading">.*?</h4>)'
        html_body = re.sub(
            rf'<p>\s*({block_alt}(?:\s*<br\s*/?>\s*{block_alt})*)\s*</p>',
            lambda m: re.sub(r'<br\s*/?>', '', m.group(1)),
            html_body, flags=re.S,
        )

        # Inject inline widths because xhtml2pdf ignores CSS nth-child
        if "LDSL" in client_name:
            html_body = re.sub(r'<th>\s*S\.?No\.?\s*</th>', '<th style="width: 8%">S.No</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Test Description\s*</th>', '<th style="width: 45%">Test Description</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Sample Type.*?</th>', '<th style="width: 30%">Sample Type (Please Tick)</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Lab Name.*?</th>', '<th style="width: 17%">Lab Name / C.C. Code</th>', html_body, flags=re.I)
        elif "Healmax" in client_name:
            html_body = re.sub(r'<th>\s*S\.?No\.?\s*</th>', '<th style="width: 5%">S.No</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Patient Name\s*</th>', '<th style="width: 14%">Patient Name</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Age\s*/\s*Sex\s*</th>', '<th style="width: 8%">Age/Sex</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Test Code\s*/\s*Name\s*</th>', '<th style="width: 20%">Test Code/Name</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Sample Type\s*</th>', '<th style="width: 10%">Sample Type</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Barcode No\s*</th>', '<th style="width: 10%">Barcode No</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Date\s*/\s*Time\s*</th>', '<th style="width: 12%">Date/Time</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Customer\s*</th>', '<th style="width: 12%">Customer</th>', html_body, flags=re.I)
            html_body = re.sub(r'<th>\s*Referral Doctor\s*</th>', '<th style="width: 9%">Referral Doctor</th>', html_body, flags=re.I)
        
        html = f"<html><head><meta charset='utf-8'/><style>{pdf_css}</style></head><body>{cover}{html_body}{footer}</body></html>"
        out = io.BytesIO()
        status = pisa.CreatePDF(io.StringIO(html), dest=out, encoding="utf-8")
        return out.getvalue() if not status.err else f"PDF Error: {status.err}"
    except Exception as e:
        return f"PDF Error: {e}"

def generate_docx(md_text: str, client_name: str) -> bytes | str:
    try:
        doc = Document()
        
        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)
            
        # Cover heading
        title = doc.add_heading(f"Extraction Report: {client_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        date_p = doc.add_paragraph(
            f"Generated: {time.strftime('%d %B %Y at %I:%M %p')}"
        )
        date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_p.runs[0].font.size = Pt(9)
        date_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()
        
        clean = re.sub(r'<\|.*?\|>', '', md_text)
        in_table = False
        table    = None
        
        for line in clean.split('\n'):
            line = line.strip()
            if not line:
                if not in_table:
                    doc.add_paragraph()
                continue
            
            if line.startswith('## '):
                in_table = False
                h = doc.add_heading(line[3:], level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
            elif line.startswith('### '):
                in_table = False
                h = doc.add_heading(line[4:], level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
            elif line.startswith('- ') or line.startswith('* '):
                in_table = False
                p = doc.add_paragraph(style='List Bullet')
                text = line[2:].strip()
                parts = re.split(r'\*\*(.+?)\*\*', text)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: run.bold = True
            elif line.startswith('|'):
                if _is_table_separator(line):
                    continue
                cells = [c.strip() for c in line.strip('|').split('|')]
                cells = [c for c in cells if c or len(cells) > 1] # allow empty cells if part of row
                if not cells:
                    continue
                if not in_table:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    hdr = table.rows[0]
                    for i, c in enumerate(cells):
                        cell = hdr.cells[i]
                        cell.text = c
                        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(c)
                        run.bold = True
                        cell._tc.get_or_add_tcPr()
                    in_table = True
                else:
                    row = table.add_row()
                    for i, c in enumerate(cells):
                        if i < len(row.cells):
                            row.cells[i].text = c
            elif line.startswith('---'):
                in_table = False
            else:
                in_table = False
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: run.bold = True
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
        note = doc.add_paragraph(
            "⚠ This extraction was auto-generated by Keppler AI from scanned case file documents. "
            "Verify all clinical data against original records before use."
        )
        note.runs[0].font.size = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        return f"Word Error: {e}"

def generate_excel(md_text: str) -> bytes | None:
    if not md_text.strip():
        return None
        
    try:
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine='openpyxl') as writer:
            # --- 1. CONSOLIDATED REPORT (Primary Sheet) ---
            report_data = []
            lines = md_text.split('\n')
            
            # Extract key-value pairs (like **Key:** Value)
            metadata = []
            other_text = []
            tables = []
            current_table = []
            
            for line in lines:
                raw_line = line.strip()
                if not raw_line: continue
                
                # Table detection — require real Markdown table syntax (leading '|'),
                # not just 2+ pipes. The LDSL/Healmax blueprints also use '|' as an
                # informal separator in plain text lines (e.g. "**Age:** 34 | **Sex:** F"),
                # which a "count('|') >= 2" check would misclassify as a table row.
                if raw_line.startswith('|'):
                    if _is_table_separator(raw_line):
                        continue
                    cells = [c.strip() for c in raw_line.split('|') if c.strip()]
                    if cells:
                        current_table.append(cells)
                    continue
                else:
                    if current_table:
                        if len(current_table) > 1:
                            tables.append(current_table)
                        current_table = []

                # Metadata detection (Bold keys) — split on '|' first so a single line
                # packing several "**Key:** Value" fields (as LDSL/Healmax headers do)
                # becomes one metadata row per field instead of one merged blob.
                if '**' in raw_line and ':' in raw_line:
                    for segment in raw_line.split('|'):
                        segment = segment.strip()
                        if not segment:
                            continue
                        parts = re.split(r'\*\*|\*', segment)
                        cleaned = " ".join([p.strip() for p in parts if p.strip()])
                        if cleaned:
                            metadata.append([cleaned])
                else:
                    other_text.append([raw_line])
            
            if current_table and len(current_table) > 1:
                tables.append(current_table)

            # Build the master list for the first sheet
            # [A] Patient / Form Metadata
            if metadata:
                report_data.append(["--- FORM DETAILS ---"])
                report_data.extend(metadata)
                report_data.append([""]) # Spacer
            
            # [B] Tables
            if tables:
                for i, table in enumerate(tables):
                    report_data.append([f"--- TABLE {i+1} ---"])
                    report_data.extend(table)
                    report_data.append([""]) # Spacer
            
            # [C] Other Text / Footer
            if other_text:
                report_data.append(["--- ADDITIONAL INFORMATION ---"])
                report_data.extend(other_text)
            
            # Create the Consolidated Sheet
            # Find max width to avoid dataframe errors
            max_cols = max([len(r) for r in report_data]) if report_data else 1
            df_cons = pd.DataFrame(report_data, columns=[f"Col {i+1}" for i in range(max_cols)])
            df_cons.to_excel(writer, index=False, header=False, sheet_name='Consolidated Report')
            
            # --- 2. CLEAN TABLE SHEETS (Individual sheets for analysis) ---
            if tables:
                for i, table in enumerate(tables):
                    header = table[0]
                    rows = table[1:]
                    cleaned_rows = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                    df_table = pd.DataFrame(cleaned_rows, columns=header)
                    sheet_name = f'Table_{i+1}'
                    df_table.to_excel(writer, index=False, sheet_name=sheet_name)
            
            # --- 3. RAW EXTRACTION SHEET ---
            paragraphs = [p.strip() for p in md_text.split('\n') if p.strip()]
            df_full = pd.DataFrame({"Raw Content": paragraphs})
            df_full.to_excel(writer, index=False, sheet_name='Raw Text')
            
        return out.getvalue()
    except Exception as e:
        print(f"Excel Error: {e}")
        # Final fallback
        try:
            out = io.BytesIO()
            pd.DataFrame({"Extraction": [md_text]}).to_excel(out, index=False)
            return out.getvalue()
        except:
            return None

def generate_json(md_text: str) -> str:
    if not md_text.strip(): return "{}"
    
    data = {"metadata": {}, "tables": [], "text": []}
    lines = md_text.split('\n')
    current_table = []
    
    for line in lines:
        raw = line.strip()
        if not raw: continue
        
        if raw.startswith('|'):
            if _is_table_separator(raw): continue
            cells = [c.strip() for c in raw.split('|') if c.strip()]
            if cells: current_table.append(cells)
            continue
        else:
            if current_table and len(current_table) > 1:
                header = current_table[0]
                rows = current_table[1:]
                table_data = []
                for row in rows:
                    row_dict = {}
                    for i, h in enumerate(header):
                        row_dict[h] = row[i] if i < len(row) else ""
                    table_data.append(row_dict)
                data["tables"].append(table_data)
                current_table = []

        if '**' in raw and ':' in raw:
            for segment in raw.split('|'):
                segment = segment.strip()
                if not segment:
                    continue
                parts = segment.split('**')
                if len(parts) >= 3:
                    key = parts[1].replace(':', '').strip()
                    val = "".join(parts[2:]).strip()
                    if key:
                        data["metadata"][key] = val
                elif segment:
                    data["text"].append(segment)
        else:
            data["text"].append(raw)
            
    if current_table and len(current_table) > 1:
        header = current_table[0]
        rows = current_table[1:]
        table_data = []
        for row in rows:
            row_dict = {}
            for i, h in enumerate(header):
                row_dict[h] = row[i] if i < len(row) else ""
            table_data.append(row_dict)
        data["tables"].append(table_data)
        
    return json.dumps(data, indent=2)

# ---------------------------------------------------------------------------
# 6.  BLUEPRINTS
# ---------------------------------------------------------------------------
BLUEPRINTS = {
    "Universal OCR (Any Text)": {
        "identity": "You are a state-of-the-art Universal OCR Vision Model, operating at the same capability level as Google Vision and OpenAI GPT-4V. You possess vast world knowledge across all domains (medical, legal, financial, etc.).",
        "structure": "Organize the extracted text into a clean, professional Markdown format. Use Markdown Headers (###) for sections, **bold text** for keys/labels, bullet points (-) for lists, and Markdown Tables (|---|) for any tabular or grid-like data.",
        "instructions": "MANDATORY: Extract ANY kind of text, across different font styles, variations, and especially messy cursive handwriting with 100% zero-shot accuracy. If the document is handwritten (e.g. a medical prescription), actively use medical domain knowledge to decipher it. Look for standard abbreviations like 'Tab', 'Cap', 'Syr'. Do not output your internal reasoning.",
        "rules": "2. EXHAUSTIVE TRANSCRIPTION: Do not stop generating until the very last word of the image is transcribed.\n3. CONTEXTUAL REASONING: Never blindly transcribe letters that form nonsensical phrases (like 'Tarsal Lateral') if a highly probable domain-specific term (like 'Tab Letoval') fits the visual strokes better. Cross-reference letter shapes against known pharmaceutical brand names.\n4. NO PREAMBLE: Start directly with the extracted data. Do not output the document domain or any greetings.\n5. FORMATTING: Use Markdown to give the text a professional structure."
    },
    "LDSL Diagnostics": {
        "identity": "You are a high-performance medical OCR engine (Qwen 2.5 VL optimized), specialized in reading LDSL Diagnostics test requisition forms.",
        "structure": """\
**Patient Name:** [Full Name as written, preserve letter spacing if boxed] | **Age:** [Age] | **Sex:** [M/F]
**Referred Doctor:** [Doctor Name] | **Collection Date:** [DD-MM-YYYY] | **Collection Time:** [HH:MM AM/PM]

| S.No | Test Description | Sample Type (Please Tick) | Lab Name / C.C. Code |
|------|-------------------|----------------------------|------------------------|
| 1    | [Test Name]       | [Ticked/circled sample type, exactly as marked] | [Code if printed, else blank] |

**History for Quadruple/Triple/Double Marker Tests:**
Date of Birth: [DOB] | Weight: [Weight]
LMP: [LMP] | Gestation Age: [Gestation Age]
Diabetes Status: [Yes/No/N/A] | Gestation: [Single/Twin]
Ultrasound Findings: [Details]

**Footer:** Collection Executive: [Name] | Date: [DD-MM-YYYY] | Checked by: [Name]""",
        "instructions": "MANDATORY: Transcribe every handwritten entry exactly. Use the 4-column Markdown table for all 10 test rows — do NOT merge the 'Sample Type' and 'Lab Name / C.C. Code' columns, they are always separate even if visually close together. Do not omit the Patient Name, History, or Footer sections.",
        "rules": "2. SPATIAL AWARENESS: Maintain the exact layout and hierarchy seen in the document — S.No, Test Description, Sample Type, and Lab Name/C.C. Code are four distinct columns, never combine them into one cell.\n3. HANDWRITING & TICKS: Transcribe every handwritten scribble or mark. If a sample type box is ticked/circled, write only that ticked option; if none is marked, leave the cell blank rather than guessing.\n4. NUMERICAL PRECISION: Do not round or alter any numbers, dates, or measurements. If a date or time is partially cut off or ambiguous, transcribe exactly what is visible rather than inferring the missing part.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary.\n6. TABLE INTEGRITY: Ensure all 10 test rows are present and every column is populated based on the visual rows, in reading order top to bottom.\n7. MISSING DATA: If any field is missing from the document, leave it completely blank. Do NOT write 'Not documented', 'N/A', or 'None' unless that literal text is printed on the form."
    },
    "Healmax Diagnostics": {
        "identity": "You are a high-performance medical OCR engine (Qwen 2.5 VL optimized).",
        "structure": """\
**Franchisee Code:** ___ | **Date:** ___

| S.No | Patient Name | Age/Sex | Test Code/Name | Sample Type | Barcode No | Date/Time | Customer | Referral Doctor |
|------|-------------|---------|----------------|-------------|------------|-----------|----------|-----------------|""",
        "instructions": "Fill all 9 table columns. Do NOT merge or skip any column.",
        "rules": "2. SPATIAL AWARENESS: Maintain the exact layout and hierarchy seen in the document.\n3. HANDWRITING: Transcribe every handwritten scribble or mark. If a checkmark is present in a box, represent it as [x].\n4. NUMERICAL PRECISION: Do not round or alter any numbers, dates, or measurements.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary.\n6. TABLE INTEGRITY: Ensure every column of the table is populated correctly based on the visual rows.\n7. MISSING DATA: If any field is missing from the document, leave it completely blank. Do NOT write 'Not documented', 'N/A', or 'None'."
    },
    "Handwritten Medical Prescription": {
        "identity": "You are an expert Clinical Pharmacist AI. You specialize in deciphering highly illegible, cursive doctor handwriting.",
        "structure": "Organize the extracted text into a clean Markdown format. Use bullet points for medications, including dosages and frequencies if present.",
        "instructions": "MANDATORY: Actively use medical domain knowledge to decipher the handwriting. Look for standard abbreviations like 'Tab' (Tablet), 'Cap' (Capsule), 'Syr' (Syrup), 'Inj' (Injection). Cross-reference letter shapes against known pharmaceutical brand names and generic drugs.",
        "rules": "2. CONTEXTUAL REASONING: Cursive shapes that look like 'Tarsal Lateral' or 'Tarsal Evaduel' in a medication list are almost certainly misreadings of 'Tab Letoval' or 'Tab Evadiol'. Use your semantic medical knowledge to read the true intent of the scribble.\n3. NO PREAMBLE: Start directly with the extracted data.\n4. FIDELITY: Do not hallucinate generic names if a brand name is written."
    }
}

# ---------------------------------------------------------------------------
# 7.  PROMPT BUILDER
# ---------------------------------------------------------------------------
def build_prompt(client: str, page_info: str = "") -> str:
    bp = BLUEPRINTS.get(client, BLUEPRINTS["Universal OCR (Any Text)"])
    page_note = f" ({page_info})" if page_info else ""
    return f"""{bp['identity']}{page_note}

Your objective: Extract EVERY piece of text from the image with 100% fidelity.

Output Format:
{bp['structure']}

Rules for 100% Accuracy:
1. MANDATORY: {bp['instructions']}
{bp['rules']}"""

# ---------------------------------------------------------------------------
# 8.  CORE PIPELINE (Streamlit-free — used by both render_ocr_app() and the FastAPI worker)
# ---------------------------------------------------------------------------
def process_single_page(raw_img: Image.Image, label: str, idx: int, total_pages: int,
                         client: str, progress_cb=None) -> dict:
    """
    Run the full single-page pipeline (layout detection -> reading order ->
    async region OCR -> table extraction -> medical correction -> confidence
    fusion) on one already-loaded page image.

    This is the unit of work the Celery OCR chord fans out per-page (see
    workers/celery_app.py) so a multi-thousand-page document is processed as
    many independent, retryable tasks across workers instead of one long
    sequential loop in a single process. `run_ocr_pipeline` below is a thin
    sequential wrapper around this for direct/synchronous callers.

    progress_cb, if given, is called with a float in [0, 1] scoped to *this
    page only* (0.0 at start, ~0.85 once correction is queued).

    Returns {"label": str, "text": str, "predictions": list[dict]}.
    """
    if progress_cb: progress_cb(0.02)

    raw_img = apply_orientation(raw_img)

    if progress_cb: progress_cb(0.05)

    page_info = f"{label} of {total_pages}" if total_pages > 1 else ""
    prompt = build_prompt(client, page_info)

    if progress_cb: progress_cb(0.15)
    
    # We use AsyncRegionOCR's robust retry logic directly on the full page image,
    # completely bypassing the destructive YOLO layout shredder.
    async_ocr = AsyncRegionOCR(max_concurrent=5, model_name="qwen2.5-vl-7b")
    
    w, h = raw_img.size
    
    # Qwen2.5-VL handles full-page table extraction natively in one shot. 
    # Bypassing the YOLO shredder completely eliminates queueing delays and drops processing time to ~5 seconds.
    if False:
        # --- PATH A: Shredded Layout Mode (Strict Tables) ---
        if progress_cb: progress_cb(0.15)
        detector = DocLayoutDetector()
        regions = detector.detect_regions(raw_img)

        if progress_cb: progress_cb(0.30)
        ro_engine = ReadingOrderEngine()
        w, h = raw_img.size
        ordered_mapping = ro_engine.reconstruct(regions, page_width=w, page_height=h)

        order_dict = {item['region_id']: item['reading_order'] for item in ordered_mapping}
        regions.sort(key=lambda r: order_dict.get(r['region_id'], 9999))

        non_table_regions = []
        table_regions = []
        for r in regions:
            if r['region_type'] == 'Table':
                table_regions.append(r)
            else:
                non_table_regions.append(r)

        if progress_cb: progress_cb(0.45)
        table_extractor = TableExtractor()

        structured_results = asyncio.run(async_ocr.process_page_regions(
            regions=non_table_regions,
            raw_img=raw_img,
            page_num=idx + 1,
            prompt=prompt
        ))

        if progress_cb: progress_cb(0.65)
        for t_reg in table_regions:
            pad = 10
            box = t_reg['bbox']
            crop_box = (max(0, box[0] - pad), max(0, box[1] - pad), min(w, box[2] + pad), min(h, box[3] + pad))
            table_img = raw_img.crop(crop_box)

            df = asyncio.run(table_extractor.extract(table_img, async_ocr, page_num=idx + 1))
            if not df.empty:
                md_table = df.to_markdown(index=False)
                structured_results.append({
                    "page": idx + 1,
                    "region_id": t_reg["region_id"],
                    "region_type": "Table",
                    "text": md_table,
                    "bbox": box,
                    "predictions": []
                })
                
        # Sort back to reading order
        structured_results.sort(key=lambda r: order_dict.get(r.get('region_id', ''), 9999))
        
        extracted_text = "\n\n".join([r['text'] for r in structured_results if r['text']])
        predictions = []
        for r in structured_results:
            predictions.extend(r.get('predictions', []))
            
    else:
        # --- PATH B: Single-Shot Mode (Handwriting Context) ---
        if progress_cb: progress_cb(0.40)
        
        extracted_text, strategy_used, predictions, ocr_confidence = asyncio.run(
            async_ocr._call_model_with_retry_async(raw_img, prompt)
        )
        
        if progress_cb: progress_cb(0.80)
    
    if progress_cb: progress_cb(0.85)
    
    page_preds = []
    if extracted_text:
        # Ground the predictions to the full page since layout detector was bypassed.
        full_page_bbox = [0, 0, w, h]
        for pred in predictions:
            # layout_confidence is always 1.0 for the full page
            sem_conf = float(pred.get("Confidence", 0.0))
            final_conf = ConfidenceEngine.calculate_final_confidence(1.0, sem_conf)
            pred["Confidence"] = f"{final_conf:.2f}"
            pred["page"] = idx + 1
            pred["bbox"] = full_page_bbox
            pred["ocr_confidence"] = ocr_confidence
            pred["ocr_model_used"] = "qwen2.5-vl-7b"
            pred["region_id"] = "full-page"
            pred["region_type"] = "Full Page"
            page_preds.append(pred)
            
        page_text = extracted_text
    else:
        page_text = f"*[{label}: Extraction failed — content too short or empty]*"

    if progress_cb: progress_cb(0.95)

    return {"label": label, "text": page_text, "predictions": page_preds}


def run_ocr_pipeline(pages: list[Image.Image], client: str, progress_cb=None) -> dict:
    """
    Sequential wrapper around process_single_page for direct/synchronous callers
    (tests, small documents processed in-process). The distributed path
    (workers/celery_app.py) calls process_single_page directly, one page per
    Celery task, instead of using this loop.

    progress_cb, if given, is called with a float in [0, 1] as pages complete.

    Returns {"pages": [(label, text), ...], "combined": str, "predictions": list[dict], "elapsed": float}.
    """
    start_time = time.time()
    total_pages = len(pages)
    page_labels = [f"Page {i+1}" for i in range(total_pages)]

    all_pages_text = []
    all_preds = []

    for idx, (raw_img, label) in enumerate(zip(pages, page_labels)):
        base_pct = idx / total_pages
        page_slice = 1.0 / total_pages

        def page_progress_cb(p, base_pct=base_pct, page_slice=page_slice):
            if progress_cb: progress_cb(base_pct + p * page_slice)

        result = process_single_page(raw_img, label, idx, total_pages, client, progress_cb=page_progress_cb)
        all_pages_text.append((result["label"], result["text"]))
        all_preds.extend(result["predictions"])

    if progress_cb:
        progress_cb(1.0)

    elapsed = time.time() - start_time

    if total_pages == 1:
        combined = all_pages_text[0][1]
    else:
        parts = [f"---\n### {lbl}\n\n{txt}" for lbl, txt in all_pages_text]
        combined = "\n\n".join(parts)

    return {
        "pages": all_pages_text,
        "combined": combined,
        "predictions": all_preds,
        "elapsed": elapsed,
    }

